import io
import asyncio
from fastapi import HTTPException
from pypdf import PdfReader
from docx import Document as DocxDocument

def _sync_extract_pdf(content: bytes) -> str:
    """Synchronous PDF extraction executed inside a thread pool."""
    try:
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse PDF document: {str(e)}"
        )

def _sync_extract_docx(content: bytes) -> str:
    """Synchronous DOCX extraction executed inside a thread pool."""
    try:
        doc = DocxDocument(io.BytesIO(content))
        text_parts = []
        
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                
        # Extract table cell text (resumes sometimes utilize tables for alignment)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in text_parts:
                        text_parts.append(cell_text)
                        
        return "\n".join(text_parts).strip()
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse DOCX document: {str(e)}"
        )

async def extract_text_async(content: bytes, filename: str) -> str:
    """
    Concurrently extracts text from a file (PDF or DOCX) by running 
    the parsing process in a background worker thread.
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".pdf"):
        return await asyncio.to_thread(_sync_extract_pdf, content)
    elif filename_lower.endswith(".docx"):
        return await asyncio.to_thread(_sync_extract_docx, content)
    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file extension. Only .pdf and .docx are supported."
        )
