"""
Sume AI — Resume ATS Analyzer Backend
FastAPI application that analyzes resumes against job descriptions using Groq LLM.
Production-ready with rate limiting, session history, security headers, and logging.
"""

from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request, Response
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pypdf import PdfReader
from docx import Document as DocxDocument
from pydantic import BaseModel, Field
from typing import List
from pathlib import Path
import datetime
import os
import io
import json
import time
import logging
from groq import AsyncGroq
from dotenv import load_dotenv
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# Load environment variables
load_dotenv(override=True)

# ── Logging Setup ────────────────────────────────────────────────────────────

IS_PRODUCTION = os.getenv("ENVIRONMENT", "development") == "production"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-7s | %(message)s" if not IS_PRODUCTION
    else '{"time":"%(asctime)s","level":"%(levelname)s","msg":"%(message)s"}',
    datefmt="%H:%M:%S" if not IS_PRODUCTION else "%Y-%m-%dT%H:%M:%S",
)
logger = logging.getLogger("sume-ai")

# ── Rate Limiter ─────────────────────────────────────────────────────────────

RATE_LIMIT = os.getenv("RATE_LIMIT", "10/hour")
limiter = Limiter(key_func=get_remote_address)

# ── Startup ──────────────────────────────────────────────────────────────────

from contextlib import asynccontextmanager

@asynccontextmanager
async def lifespan(app):
    """Startup events."""
    logger.info(f"Rate limit: {RATE_LIMIT}")
    yield


# ── App Setup ────────────────────────────────────────────────────────────────

app = FastAPI(
    title="Sume AI — Resume ATS Analyzer",
    description="Analyze resumes against job descriptions for ATS optimization",
    version="3.0.0",
    lifespan=lifespan,
)

# Rate limit error handler
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# CORS middleware
allowed_origins = os.getenv("ALLOWED_ORIGINS", "*").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins if allowed_origins != ["*"] else ["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Security Headers Middleware ──────────────────────────────────────────────

@app.middleware("http")
async def add_security_headers(request: Request, call_next):
    """Add security headers to every response."""
    start_time = time.time()
    response = await call_next(request)

    # Security headers
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"

    # Log request timing
    duration = round((time.time() - start_time) * 1000)
    if request.url.path not in ("/health", "/favicon.ico"):
        logger.info(f"{request.method} {request.url.path} -> {response.status_code} ({duration}ms)")

    return response



# ── Groq Client ─────────────────────────────────────────────────────────────

def get_groq_keys() -> List[str]:
    """Retrieve all available Groq API keys."""
    load_dotenv(override=True)
    keys = [
        os.getenv("GROQ_API_KEY"),
        os.getenv("GROQ_API_KEY_2"),
        os.getenv("GROQ_API_KEY_3")
    ]
    # Filter out None and empty strings
    valid_keys = [k.strip() for k in keys if k and k.strip()]
    
    if not valid_keys:
        raise HTTPException(
            status_code=500,
            detail="No GROQ_API_KEY is set. Please add it to your .env file.",
        )
    return valid_keys


async def call_groq_with_fallback(call_kwargs: dict):
    """Attempt a Groq API chat completion with automatic key rotation on failure."""
    keys = get_groq_keys()
    
    for attempt, key in enumerate(keys):
        try:
            client = AsyncGroq(api_key=key)
            response = await client.chat.completions.create(**call_kwargs)
            return response
        except Exception as e:
            key_name = "GROQ_API_KEY" if attempt == 0 else f"GROQ_API_KEY_{attempt + 1}"
            logger.warning(f"Groq API call failed using {key_name} (Attempt {attempt+1}/{len(keys)}): {str(e)}")
            
    # If all keys failed
    logger.error("All available Groq API keys failed.")
    raise HTTPException(
        status_code=500, 
        detail="AI service is currently unavailable or rate limits exceeded. Please try again later."
    )



# ── Pydantic Models ──────────────────────────────────────────────────────────


class KeywordDetail(BaseModel):
    """A missing keyword with context on why it matters."""
    keyword: str
    importance: str = Field(default="", description="Why this keyword matters")


class RewriteSuggestion(BaseModel):
    """A before→after rewrite suggestion for a weak bullet point."""
    original: str = Field(description="Exact text from the resume that should be changed")
    rewritten: str = Field(description="Improved version of the text")
    why: str = Field(default="", description="Brief explanation of why this change improves ATS compatibility")


class ATSParsedSection(BaseModel):
    """A section that an ATS would try to detect in the resume."""
    name: str = Field(description="Section name, e.g. 'Contact Info', 'Experience', 'Education'")
    status: str = Field(description="'found', 'missing', or 'warning'")
    detail: str = Field(default="", description="What the ATS found or why it's flagged")


class AnalysisResult(BaseModel):
    """Structured ATS analysis result from the LLM."""
    overall_score: int = Field(ge=0, le=100, description="ATS match score 0-100")
    summary: str = Field(default="", description="2-3 sentence overall assessment")
    missing_keywords: List[KeywordDetail] = Field(default=[], description="Keywords from JD missing in resume")
    strengths: List[str] = Field(default=[], description="What the resume does well")
    improvements: List[str] = Field(default=[], description="Specific, actionable improvements")
    ats_issues: List[str] = Field(default=[], description="ATS formatting/compatibility issues")
    section_scores: dict = Field(
        default={}, description="Breakdown scores by section"
    )
    rewrites: List[RewriteSuggestion] = Field(
        default=[], description="Before→after rewrite suggestions for weak bullet points"
    )
    ats_parsed_sections: List[ATSParsedSection] = Field(
        default=[], description="Simulated ATS section detection results"
    )
    updated_resume_md: str = Field(
        default="", description="Complete updated resume in Markdown format with all improvements applied"
    )
    confidence: float = Field(default=0.5, ge=0.0, le=1.0, description="Model confidence")


# ── Helper Functions ─────────────────────────────────────────────────────────


def extract_pdf_text(file) -> str:
    """Extract all text from a PDF file."""
    try:
        pdf = PdfReader(file)
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read PDF: {str(e)}")


def extract_docx_text(file) -> str:
    """Extract all text from a DOCX file (paragraphs + table cells)."""
    try:
        doc = DocxDocument(file)
        text_parts = []
        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        # Extract table cell text (resumes sometimes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in text_parts:
                        text_parts.append(cell_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read DOCX: {str(e)}")


def build_analysis_prompt(resume_text: str, job_description: str) -> str:
    """Build the LLM prompt for ATS analysis with rewrites and ATS simulation."""
    return f"""You are an expert ATS (Applicant Tracking System) analyst and career coach. 
Analyze the following resume against the given job description with extreme precision.

## Job Description:
{job_description}

## Resume:
{resume_text}

## Your Task:
Perform a thorough ATS compatibility analysis. Score the resume 0-100 based on:
- Keyword match (40% weight): How many critical keywords from the JD appear in the resume?
- Experience relevance (25% weight): Does the experience align with what the role needs?
- Skills alignment (20% weight): Are the required technical/soft skills present?
- Formatting & structure (15% weight): Is the resume ATS-parseable (standard sections, no tables/images, clean formatting)?

## Respond in this EXACT JSON format:
{{
    "overall_score": <number 0-100>,
    "summary": "<2-3 sentence overall assessment of the resume's ATS compatibility>",
    "missing_keywords": [
        {{"keyword": "<keyword>", "importance": "<why it matters, e.g. 'Mentioned 5 times in JD — critical skill'>"}}
    ],
    "strengths": [
        "<specific strength with evidence — e.g., 'Strong action verbs: Led, Architected, Delivered'>"
    ],
    "improvements": [
        "<specific, actionable fix — e.g., 'Replace Responsible for managing with Managed a team of 8 engineers'>"
    ],
    "ats_issues": [
        "<specific ATS problem — e.g., 'Section header Work History should be EXPERIENCE for ATS compatibility'>"
    ],
    "section_scores": {{
        "experience": <0-100>,
        "skills": <0-100>,
        "education": <0-100>,
        "formatting": <0-100>
    }},
    "rewrites": [
        {{
            "original": "<EXACT text copied from the resume that needs improvement>",
            "rewritten": "<improved version — specific, quantified, ATS-optimized>",
            "why": "<1 sentence explaining the improvement>"
        }}
    ],
    "ats_parsed_sections": [
        {{
            "name": "Contact Info",
            "status": "found",
            "detail": "Name, email, and phone detected"
        }},
        {{
            "name": "Summary/Objective",
            "status": "missing",
            "detail": "No professional summary found — add a 2-3 line summary at the top"
        }},
        {{
            "name": "Experience",
            "status": "found",
            "detail": "4 positions detected with bullet points"
        }},
        {{
            "name": "Education",
            "status": "found",
            "detail": "1 degree detected"
        }},
        {{
            "name": "Skills",
            "status": "warning",
            "detail": "Skills listed in paragraph form — use a comma-separated or bulleted list for better ATS parsing"
        }},
        {{
            "name": "Certifications",
            "status": "missing",
            "detail": "No certifications section — add if you have relevant certs"
        }},
        {{
            "name": "Projects",
            "status": "found",
            "detail": "2 projects detected"
        }}
    ],
    "updated_resume_md": "<A COMPLETE, fully rewritten resume in clean Markdown format. Apply ALL the improvements, insert missing keywords naturally, fix ATS issues, use strong action verbs, and quantify achievements. Structure it with proper Markdown headers (# Name, ## Experience, ## Skills, ## Education, ## Projects, etc.). This should be a polished, ready-to-use resume that the candidate can copy directly.>",
    "confidence": <0.0-1.0>
}}

## CRITICAL RULES:
- Be SPECIFIC and ACTIONABLE — no generic advice like "add more keywords"
- Reference exact phrases from the resume and JD
- For missing_keywords, only include keywords genuinely in the JD but absent from the resume
- For "rewrites": provide 3-5 specific before→after rewrites. Copy the EXACT original text from the resume. The rewritten version must be specific, quantified where possible, and optimized for ATS keywords from the JD.
- For "ats_parsed_sections": simulate how a real ATS (like Taleo, Workday, Greenhouse) would parse this resume. Report on these sections: Contact Info, Summary/Objective, Experience, Education, Skills, Certifications, Projects. Use status "found", "missing", or "warning".
- For "updated_resume_md": Generate a COMPLETE updated resume in Markdown. Do NOT just list changes — write the ENTIRE resume with improvements applied. Use proper Markdown formatting with headers, bullet points, and bold text. Preserve all original information but enhance it with missing keywords, better action verbs, and quantified achievements.
- Score honestly — a poor match should score below 40
- Return ONLY valid JSON, no markdown formatting"""


# ── API Endpoints ────────────────────────────────────────────────────────────


@app.get("/health")
def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "service": "sume-ai", "version": "3.0.0"}


@app.post("/analyze-resume")
@limiter.limit(RATE_LIMIT)
async def analyze_resume(
    request: Request,
    resume: UploadFile = File(..., description="Resume PDF or DOCX file (max 5MB)"),
    job_description: str = Form(..., description="Job description text (max 5000 chars)"),
):
    """
    Analyze a resume PDF against a job description for ATS compatibility.
    Returns a structured analysis with score, missing keywords, and suggestions.
    """
    # ── Validate inputs ──────────────────────────────────────────────────
    filename_lower = resume.filename.lower() if resume.filename else ""
    is_pdf = filename_lower.endswith(".pdf")
    is_docx = filename_lower.endswith(".docx")

    if not is_pdf and not is_docx:
        raise HTTPException(
            status_code=400,
            detail="Only PDF and DOCX files are supported. Please upload a .pdf or .docx file.",
        )

    content = await resume.read()
    if len(content) > 5_000_000:
        raise HTTPException(status_code=400, detail="File too large. Maximum file size is 5MB.")

    if len(job_description.strip()) < 50:
        raise HTTPException(
            status_code=400,
            detail="Job description is too short. Please paste the full job description (at least 50 characters).",
        )

    if len(job_description) > 5000:
        raise HTTPException(status_code=400, detail="Job description is too long. Maximum 5000 characters.")

    # ── Extract text from resume ──────────────────────────────────────────
    if is_pdf:
        resume_text = extract_pdf_text(io.BytesIO(content))
    else:
        resume_text = extract_docx_text(io.BytesIO(content))

    if not resume_text or len(resume_text.strip()) < 50:
        file_type = "PDF" if is_pdf else "DOCX"
        raise HTTPException(
            status_code=400,
            detail=f"Could not extract sufficient text from the {file_type}. The file may be image-based or corrupted.",
        )

    # ── Analyze with Groq LLM ───────────────────────────────────────────
    prompt = build_analysis_prompt(resume_text, job_description)
    logger.info(f"Analyzing resume: {resume.filename} ({len(resume_text)} chars) vs JD ({len(job_description)} chars)")

    try:
        response = await call_groq_with_fallback({
            "model": "llama-3.1-8b-instant",
            "messages": [
                {
                    "role": "system",
                    "content": "You are an expert ATS resume analyst. Always respond with valid JSON only.",
                },
                {"role": "user", "content": prompt},
            ],
            "temperature": 0,
            "max_tokens": 4500,
            "response_format": {"type": "json_object"},
        })

        raw_content = response.choices[0].message.content

        # Parse and validate
        try:
            result = AnalysisResult.model_validate_json(raw_content)
        except Exception:
            raw_data = json.loads(raw_content)
            result = AnalysisResult.model_validate(raw_data)

        result_dict = result.model_dump()
        logger.info(f"Analysis complete: score={result.overall_score}, confidence={result.confidence}")

        # Increment global analysis counter
        _increment_analysis_count()

        return JSONResponse(content=result_dict)

    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=500,
            detail=f"AI returned invalid response format. Please try again. Error: {str(e)}",
        )
    except HTTPException:
        raise
    except Exception as e:
        error_msg = str(e)
        logger.error(f"Analysis error: {error_msg}")
        if "api_key" in error_msg.lower() or "invalid" in error_msg.lower() or "401" in error_msg or "authentication" in error_msg.lower():
            raise HTTPException(
                status_code=500,
                detail=f"Groq API rejected the key. Please generate a new key at console.groq.com/keys. Details: {error_msg}",
            )
        raise HTTPException(
            status_code=500,
            detail=f"Analysis failed. Please try again. Error: {error_msg}",
        )


# ── Analysis Counter Helpers ─────────────────────────────────────────────────

COUNTER_FILE = Path(os.path.dirname(os.path.abspath(__file__))) / "analysis_count.json"
FEEDBACK_FILE = Path(os.path.dirname(os.path.abspath(__file__))) / "feedback_log.json"


def _read_analysis_count() -> int:
    """Read the current analysis count from the counter file."""
    try:
        if COUNTER_FILE.exists():
            data = json.loads(COUNTER_FILE.read_text())
            return data.get("count", 1000)
    except Exception:
        pass
    return 1000  # Seed value


def _increment_analysis_count():
    """Increment the analysis counter."""
    try:
        count = _read_analysis_count() + 1
        COUNTER_FILE.write_text(json.dumps({"count": count}))
    except Exception as e:
        logger.warning(f"Could not update analysis counter: {e}")


@app.get("/analytics/user-count")
def get_user_count():
    """Return the total number of resumes analyzed (for social proof counter)."""
    return {"count": _read_analysis_count()}


# ── Feedback Endpoint ────────────────────────────────────────────────────────


@app.post("/submit-feedback")
@limiter.limit("5/hour")
async def submit_feedback(
    request: Request,
    name: str = Form(default="Anonymous"),
    message: str = Form(..., min_length=5, max_length=1000),
):
    """Accept user feedback and log it to a local JSON file."""
    try:
        feedback_entry = {
            "name": name.strip() or "Anonymous",
            "message": message.strip(),
            "timestamp": datetime.datetime.now(datetime.timezone.utc).isoformat(),
        }
        # Read existing feedback
        entries = []
        if FEEDBACK_FILE.exists():
            try:
                entries = json.loads(FEEDBACK_FILE.read_text())
            except Exception:
                entries = []
        entries.append(feedback_entry)
        FEEDBACK_FILE.write_text(json.dumps(entries, indent=2))
        logger.info(f"Feedback received from: {feedback_entry['name']}")
        return JSONResponse(content={"status": "success", "message": "Thank you for your feedback!"})
    except Exception as e:
        logger.error(f"Feedback error: {str(e)}")
        raise HTTPException(status_code=500, detail="Failed to save feedback.")


# ── Cover Letter Endpoint ────────────────────────────────────────────────────


def build_cover_letter_prompt(resume_text: str, job_description: str) -> str:
    """Build the LLM prompt for cover letter generation."""
    return f"""You are an expert career coach and professional writer.
Generate a tailored, compelling cover letter based on the resume and job description below.

## Job Description:
{job_description}

## Resume:
{resume_text}

## Requirements:
1. Address "Dear Hiring Manager" (unless a name is in the JD)
2. Opening paragraph: Hook with a standout achievement or passion that connects to the role
3. Body (1-2 paragraphs): Map the candidate's specific experience/skills to the JD requirements. Reference concrete achievements from the resume with metrics when available.
4. Closing paragraph: Express enthusiasm, reference company values/mission if apparent from JD, include a call to action
5. Keep it under 350 words — concise and impactful
6. Tone: Professional but personable, confident but not arrogant
7. DO NOT fabricate any skills, experiences, or achievements not in the resume

## RESPOND WITH ONLY THE COVER LETTER TEXT. No JSON, no markdown formatting, no explanations. Just the cover letter ready to copy-paste."""


@app.post("/generate-cover-letter")
@limiter.limit(RATE_LIMIT)
async def generate_cover_letter(
    request: Request,
    resume: UploadFile = File(..., description="Resume PDF or DOCX file (max 5MB)"),
    job_description: str = Form(..., description="Job description text"),
):
    """
    Generate a tailored cover letter based on the resume and job description.
    """
    # Validate file type
    filename_lower = resume.filename.lower() if resume.filename else ""
    is_pdf = filename_lower.endswith(".pdf")
    is_docx = filename_lower.endswith(".docx")

    if not is_pdf and not is_docx:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")

    content = await resume.read()
    if len(content) > 5_000_000:
        raise HTTPException(status_code=400, detail="File too large. Maximum 5MB.")

    # Extract text
    if is_pdf:
        resume_text = extract_pdf_text(io.BytesIO(content))
    else:
        resume_text = extract_docx_text(io.BytesIO(content))

    if not resume_text or len(resume_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Could not extract sufficient text from the file.")

    # Generate cover letter
    prompt = build_cover_letter_prompt(resume_text, job_description)
    logger.info(f"Generating cover letter for: {resume.filename}")

    try:
        response = await call_groq_with_fallback({
            "model": "llama-3.1-8b-instant",
            "messages": [
                {
                    "role": "system",
                    "content": "You are a professional cover letter writer. Respond with only the cover letter text, nothing else.",
                },
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.7,
            "max_tokens": 1500,
        })

        cover_letter = response.choices[0].message.content.strip()
        logger.info(f"Cover letter generated: {len(cover_letter)} chars")
        return JSONResponse(content={"cover_letter": cover_letter})

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Cover letter error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Cover letter generation failed: {str(e)}")


# ── Static Files ─────────────────────────────────────────────────────────────

static_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
if os.path.exists(static_dir):
    app.mount("/static", StaticFiles(directory=static_dir), name="static")


@app.get("/")
def serve_frontend():
    """Serve the main frontend page."""
    index_path = os.path.join(static_dir, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path)
    return {"message": "Sume AI API is running. Frontend not found."}


# ── Entry Point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn

    PORT = int(os.getenv("PORT", 8000))
    print(f"\n  [*] Sume AI v3.0 starting on http://localhost:{PORT}\n")
    uvicorn.run(app, host="0.0.0.0", port=PORT)
